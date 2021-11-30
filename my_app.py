from docx import Document

doc = Document()


name = input("Enter You're Name : ")
phone = input("Enter You're Mobile Number : ")
email = input("Enter You're Email : ")

doc.add_paragraph(name + '\n' + phone + '\n' + email)

# About me section

doc.add_heading("About me", 0)
doc.add_paragraph(
    input("Tell us About yourself? ")
)

# Skills section

doc.add_heading("Skills", 0)
skill = input("Enter You're Skills : ")
p = doc.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    more_skills = input(
        "Do you have more Skills?")
    if more_skills == 'yes':
        skill = input("Enter skill :")
        p = doc.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break


doc.save('resume.docx')
